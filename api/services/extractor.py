import io

import docx
import fitz
import openpyxl


class UnsupportedFormatError(Exception):
    """Raised when the uploaded file format is not PDF, DOCX, or XLSX."""


def extract_text(content: bytes, filename: str) -> str:
    """
    Description: Extract plain text from a PDF, DOCX, or XLSX file.

    Args:
        content: Raw file bytes read from the upload
        filename: Original filename used to determine the format

    Returns:
        Extracted text as a single concatenated string
    """
    name = filename.lower()
    if name.endswith(".pdf"):
        return _extract_pdf(content)
    if name.endswith(".docx"):
        return _extract_docx(content)
    if name.endswith(".xlsx"):
        return _extract_xlsx(content)
    raise UnsupportedFormatError(filename)


def _extract_pdf(content: bytes) -> str:
    """
    Description: Extract text from PDF bytes page by page.

    Args:
        content: Raw PDF bytes

    Returns:
        Concatenated text from all pages
    """
    doc = fitz.open(stream=content, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _extract_docx(content: bytes) -> str:
    """
    Description: Extract text from DOCX bytes including table cells.

    Args:
        content: Raw DOCX bytes

    Returns:
        Concatenated text from paragraphs and table cells
    """
    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _extract_xlsx(content: bytes) -> str:
    """
    Description: Extract text from XLSX bytes cell by cell.

    Args:
        content: Raw XLSX bytes

    Returns:
        Concatenated cell values row by row across all sheets
    """
    wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            parts.extend(str(cell) for cell in row if cell is not None)
    return "\n".join(parts)
